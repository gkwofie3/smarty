import random
from datetime import datetime, timedelta
import os
import json
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.enum.style import WD_STYLE_TYPE
# New import for Pandoc conversion
import pypandoc 


# --- PATH CONFIGURATION ---
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) 
STATIC_ROOT = os.path.join(PROJECT_ROOT, 'static')
MEDIA_ROOT = os.path.join(PROJECT_ROOT, 'media')

# --- FILE PATHS ---
LOGO_PATH = os.path.join(STATIC_ROOT, "assets/img/logo/smartheader.png")
LOGO = os.path.join(STATIC_ROOT, "assets/img/logo/logo.png")
REPORT_SAVE_DIR = os.path.join(MEDIA_ROOT, "reports/fire_system")

# Define both DOCX and PDF filenames using the same timestamp
REPORT_TIMESTAMP = datetime.now().strftime('%Y%m%d_%H%M%S')
OUTPUT_DOCX = os.path.join(REPORT_SAVE_DIR, f"fire_system_report_{REPORT_TIMESTAMP}.docx")
OUTPUT_PDF = os.path.join(REPORT_SAVE_DIR, f"fire_system_report_{REPORT_TIMESTAMP}.pdf") # NEW PDF PATH
OUTPUT_JSON = os.path.join(REPORT_SAVE_DIR, "fire_report_snapshot.json")

# Deterministic random seed for reproducible sample data
random.seed(20251129)

# --- Configuration & Data Generation ---
REPORT_PERIOD_START = datetime(2025,10,1,0,0,0)
REPORT_PERIOD_END = datetime(2025,10,31,23,59,59)
GENERATED_ON = datetime(2025,11,29,8,16,32) 

NUM_ZONES = 20
NUM_LOOPS = 6
DEVICES_PER_LOOP = [random.randint(60,80) for _ in range(NUM_LOOPS)]

PANEL_PARAMS = [
    "General Alarm","General Fault","General Disablement","General Test","Sounder ON",
    "Sounder Delay","Buzzer Muted","Mains Fault","Battery Fault","Charger Fault",
    "Earth Fault","Fuse Fault","Network Fault"
]

def generate_devices(loop_index, count):
    devices = []
    for addr in range(1, count+1):
        occupied = random.random() > 0.01
        alarm = random.random() < 0.02
        fault = random.random() < 0.05
        disabled = random.random() < 0.03
        devices.append({
            "loop": loop_index+1,
            "address": addr,
            "occupied": occupied,
            "alarm": alarm,
            "fault": fault,
            "disabled": disabled
        })
    return devices

loops = []
for li in range(NUM_LOOPS):
    count = DEVICES_PER_LOOP[li]
    loops.append({
        "loop_id": li+1,
        "device_count": count,
        "devices": generate_devices(li, count)
    })

zones = []
for zi in range(NUM_ZONES):
    zone_id = zi+1
    primary_loop = random.randint(1, NUM_LOOPS)
    zones.append({
        "zone_id": zone_id,
        "description": f"Area {zone_id} (Example)",
        "alarm_count": 0,
        "fault_count": 0,
        "disablement_count": 0,
        "primary_loop": primary_loop,
        "alarm_group": "Group A" if zone_id % 2 == 0 else "Group B",
        "last_test": None
    })

device_events = []
for loop in loops:
    for dev in loop["devices"]:
        zone = random.choice(zones)
        if dev["alarm"]:
            zone["alarm_count"] += 1
        if dev["fault"]:
            zone["fault_count"] += 1
        if dev["disabled"]:
            zone["disablement_count"] += 1
        
        random_day = random.randint(0, (REPORT_PERIOD_END - REPORT_PERIOD_START).days)
        random_time = timedelta(days=random_day, hours=random.randint(0,23), minutes=random.randint(0,59))
        timestamp = (REPORT_PERIOD_START + random_time).isoformat()
        
        device_events.append({
            "timestamp": timestamp,
            "loop": dev["loop"],
            "address": dev["address"],
            "zone_id": zone["zone_id"],
            "occupied": dev["occupied"],
            "alarm": dev["alarm"],
            "fault": dev["fault"],
            "disabled": dev["disabled"]
        })

total_alarms = sum(1 for d in device_events if d["alarm"])
total_faults = sum(1 for d in device_events if d["fault"])
total_disabled = sum(1 for d in device_events if d["disabled"])

notable_events = [
    {"start":"2025-10-05T09:00:00","end":"2025-10-05T09:30:00","type":"General Test","details":"Full system monthly test","ack":"Engineer 1"},
    {"start":"2025-10-10T11:45:22","end":"2025-10-10T11:50:44","type":"Network Fault","details":"Communication loss to Loop 2","ack":"Operator 2"},
    {"start":"2025-10-15T10:00:00","end":"2025-10-15T12:00:00","type":"Zone Disablement","details":"Zone 5 disabled for HVAC maintenance","ack":"Engineer 3"},
    {"start":"2025-10-25T03:15:00","end":"2025-10-25T03:16:45","type":"General Alarm","details":"Zone 1 - Device 010 (Smoke)","ack":"Operator 1"},
    {"start":"2025-10-27T14:30:11","end":"2025-10-27T14:32:00","type":"General Alarm","details":"Zone 1 - Device 012 (Heat)","ack":"Operator 4"},
    {"start":"2025-10-30T09:11:05","end":"2025-10-30T09:12:00","type":"Network Fault","details":"Communication loss to Loop 1","ack":"Operator 2"},
    {"start":"2025-10-31T19:55:10","end":"2025-10-31T19:56:40","type":"General Alarm","details":"Zone 5 - Device 201 (Call Point)","ack":"Security A"},
]

# --- DOCX Composition Helper Functions ---
def set_document_font(document, font_name="Times New Roman"):
    """Set the default font for Normal style and common Heading styles."""
    styles = document.styles
    
    if 'Normal' in styles:
        normal_style = styles['Normal']
        font = normal_style.font
        font.name = font_name
        font.size = Pt(11)

    for i in range(1, 4): 
        style_name = f'Heading {i}'
        try:
            heading_style = styles[style_name]
            heading_font = heading_style.font
            heading_font.name = font_name
        except KeyError:
            pass

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(paragraph):
    """Add a page number field to a paragraph."""
    p = paragraph._p
    fldSimple = create_element('w:fldSimple')
    create_attribute(fldSimple, 'w:instr', 'PAGE \\* MERGEFORMAT')
    run = create_element('w:r')
    t = create_element('w:t')
    t.text = "1" 
    run.append(t)
    p.append(fldSimple)
    
def add_toc(document):
    """Add a table of contents field (TOC) to the document."""
    paragraph = document.add_paragraph()
    p = paragraph._p
    fldChar_start = create_element('w:fldChar')
    create_attribute(fldChar_start, 'w:fldCharType', 'begin')
    p.append(fldChar_start)

    run = create_element('w:r')
    instrText = create_element('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u' 
    run.append(instrText)
    p.append(run)

    fldChar_end = create_element('w:fldChar')
    create_attribute(fldChar_end, 'w:fldCharType', 'end')
    p.append(fldChar_end)

    paragraph.style = 'TOCHeading' 

# --- Function Start ---

def generate_fire_system_report():
    """Generates a comprehensive fire system report in DOCX, converts to PDF using Pandoc, and saves a data snapshot in JSON."""

    # Ensure the media directory structure exists
    if not os.path.exists(REPORT_SAVE_DIR):
        os.makedirs(REPORT_SAVE_DIR)
        print(f"Created report directory: {REPORT_SAVE_DIR}")

    doc = Document()

    # Set default font to Times New Roman
    set_document_font(doc, "Times New Roman")

    # --- HEADER (LOGO) ---
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.paragraph_format.left_indent = -Inches(0.8)
    # Add logo using the adjusted LOGO_PATH
    try:
        if os.path.exists(LOGO_PATH):
            header_p.add_run().add_picture(LOGO_PATH, width=Inches(1)) 
            header_p.add_run() 
        else:
            header_p.add_run('[Logo Placeholder: File not found at ' + LOGO_PATH + ']')
    except Exception as e:
        header_p.add_run(f'[Logo Error: {e}]') 
        
    header_p.add_run('\t \tFire System Report')
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- FOOTER (System generated fire system report + Page Number) ---
    footer = section.footer
    footer_p = footer.paragraphs[0]
    # Using GENERATED_ON for consistency
    footer_p.text = f'System generated fire system report - {GENERATED_ON.strftime("%B %d, %Y %H:%M:%S")}'

    # Add page number to the right
    footer_p.add_run('\t\t\t\tPage ') 
    add_page_number(footer_p)

    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- COVER PAGE ---
    doc.add_paragraph("\n") 
    branding_p = doc.add_paragraph()
    branding_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = branding_p.add_run("ROVID SMART TECHNOLOGIES")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = docx.shared.RGBColor(83, 2, 150)
    
    doc.add_paragraph("\n")
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_picture(LOGO, width=Inches(4))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    smarty_p = doc.add_paragraph()
    smarty_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = smarty_p.add_run("SMARTY BUILDING MANAGEMENT SYSTEM")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = docx.shared.RGBColor(179, 64, 7)
    doc.add_paragraph("\n")
    twi_p = doc.add_paragraph()
    twi_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = twi_p.add_run("TWELLIUM INDUSTRIAL COMPANY LIMITED-ACCRA")
    
    report_title_p = doc.add_paragraph("FIRE SYSTEM REPORT — LAST MONTH")
    report_title_p.alignment= WD_ALIGN_PARAGRAPH.CENTER
    
    # Format the TWELLIUM run (same formatting as the title run in your original code)
    run.font.size = Pt(20)
    run.font.bold = True
    doc.add_page_break()

    # --- TABLE OF CONTENT ---
    doc.add_heading('Table of Contents', level=1)
    add_toc(doc)
    doc.add_page_break()

    # Title
    doc.add_heading('FIRE SYSTEM REPORT — LAST MONTH', 0)
    doc.add_paragraph(f'Reporting Period: {REPORT_PERIOD_START.strftime("%B %d, %Y")} — {REPORT_PERIOD_END.strftime("%B %d, %Y")}')
    doc.add_paragraph(f'Report Generated: {GENERATED_ON.strftime("%B %d, %Y %H:%M:%S")}')
    doc.add_paragraph('System Status: ONLINE')

    # 1. EXECUTIVE SUMMARY
    doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
    doc.add_paragraph('This report summarizes fire system performance for the reporting period. It includes panel status, zone summaries, loop and device-level events, a critical event log, analysis, and recommendations.')

    # KPI table
    kpi_table = doc.add_table(rows=1, cols=3)
    kpi_table.style = 'Table Grid'
    hdr_cells = kpi_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Count'
    hdr_cells[2].text = 'Notes'

    kpis = [
        ('Total Fire Alarms', str(total_alarms), 'All alarms acknowledged & cleared'),
        ('Total General Faults', str(total_faults), 'Includes device and network faults'),
        ('Total Disablements', str(total_disabled), 'Scheduled maintenance disablements included'),
        ('Panel Downtime', '2 hrs 15 mins', 'Scheduled maintenance on Oct 15'),
        ('Zones with Activity', str(sum(1 for z in zones if z["alarm_count"] or z["fault_count"] or z["disablement_count"])), 'Out of 20 monitored zones')
    ]

    for k, v, n in kpis:
        row_cells = kpi_table.add_row().cells
        row_cells[0].text = k
        row_cells[1].text = v
        row_cells[2].text = n

    # 2. GENERAL PANEL STATUS HISTORY
    doc.add_heading('2. GENERAL PANEL STATUS HISTORY', level=1)
    doc.add_paragraph('Panel parameters tracked during the reporting period: ' + ', '.join(PANEL_PARAMS))

    panel_table = doc.add_table(rows=1, cols=4)
    panel_table.style = 'Table Grid'
    hdr = panel_table.rows[0].cells
    hdr[0].text = 'Status Parameter'; hdr[1].text = 'Total Events'; hdr[2].text = 'Longest Duration'; hdr[3].text = 'Notes'

    panel_counts = {p: random.randint(0,12) for p in PANEL_PARAMS}
    panel_counts['General Alarm'] = total_alarms
    panel_counts['General Fault'] = total_faults
    panel_counts['General Disablement'] = total_disabled
    panel_counts['Network Fault'] = panel_counts.get('Network Fault', 0)

    for p in PANEL_PARAMS:
        row = panel_table.add_row().cells
        row[0].text = p
        row[1].text = str(panel_counts.get(p,0))
        row[2].text = '00:05:22' if p == 'Network Fault' else ('00:01:49' if p == 'General Alarm' else 'N/A')
        row[3].text = 'See event log' if panel_counts.get(p,0) else ''

    # 3. ZONE STATUS SUMMARY
    doc.add_heading('3. ZONE STATUS SUMMARY', level=1)
    zone_table = doc.add_table(rows=1, cols=6)
    zone_table.style = 'Table Grid'
    zhdr = zone_table.rows[0].cells
    zhdr[0].text = 'Zone'; zhdr[1].text = 'Description'; zhdr[2].text = 'Alarms';zhdr[3].text='Faults';zhdr[4].text='Disablements';zhdr[5].text='Primary Loop'
    for z in zones:
        r = zone_table.add_row().cells
        r[0].text = str(z['zone_id'])
        r[1].text = z['description']
        r[2].text = str(z['alarm_count'])
        r[3].text = str(z['fault_count'])
        r[4].text = str(z['disablement_count'])
        r[5].text = str(z['primary_loop'])

    # 4. CRITICAL EVENT LOG
    doc.add_heading('4. CRITICAL EVENT LOG', level=1)
    elog_table = doc.add_table(rows=1, cols=6)
    elog_table.style = 'Table Grid'
    eh = elog_table.rows[0].cells
    eh[0].text='Start Time'; eh[1].text='End Time'; eh[2].text='Event Type'; eh[3].text='Details'; eh[4].text='Acknowledged By'; eh[5].text='Duration'
    for ev in notable_events:
        r = elog_table.add_row().cells
        r[0].text = ev['start']
        r[1].text = ev['end']
        r[2].text = ev['type']
        r[3].text = ev['details']
        r[4].text = ev['ack']
        # compute duration
        try:
            dstart = datetime.fromisoformat(ev['start'])
            dend = datetime.fromisoformat(ev['end'])
            dur = dend - dstart
            r[5].text = str(dur)
        except Exception:
            r[5].text = ''

    # 5. LOOP PERFORMANCE SUMMARY
    doc.add_heading('5. LOOP PERFORMANCE SUMMARY', level=1)
    loop_table = doc.add_table(rows=1, cols=5)
    loop_table.style = 'Table Grid'
    lh = loop_table.rows[0].cells
    lh[0].text='Loop'; lh[1].text='Devices Present'; lh[2].text='Occupied'; lh[3].text='Faults'; lh[4].text='Disabled'
    for loop in loops:
        total = loop['device_count']
        occupied = sum(1 for d in loop['devices'] if d['occupied'])
        faults = sum(1 for d in loop['devices'] if d['fault'])
        disabled = sum(1 for d in loop['devices'] if d['disabled'])
        r = loop_table.add_row().cells
        r[0].text = str(loop['loop_id'])
        r[1].text = str(total)
        r[2].text = str(occupied)
        r[3].text = str(faults)
        r[4].text = str(disabled)

    # 6. DEVICE-LEVEL ACTIVITY (Sample)
    doc.add_heading('6. DEVICE-LEVEL ACTIVITY (Sample)', level=1)
    doc.add_paragraph('This section lists sample device events (alarms, faults, disabled) observed during the reporting period. Only a subset of devices is shown for brevity.')

    d_table = doc.add_table(rows=1, cols=7)
    d_table.style = 'Table Grid'
    dh = d_table.rows[0].cells
    dh[0].text='Timestamp'; dh[1].text='Loop'; dh[2].text='Address'; dh[3].text='Zone'; dh[4].text='Occupied'; dh[5].text='Alarm/Fault/Disabled'; dh[6].text='Notes'
    # show up to 60 sample device events (deterministic slice)
    for d in device_events[:60]:
        r = d_table.add_row().cells
        r[0].text = d['timestamp']
        r[1].text = str(d['loop'])
        r[2].text = str(d['address'])
        r[3].text = str(d['zone_id'])
        r[4].text = 'Yes' if d['occupied'] else 'No'
        flags = []
        if d['alarm']: flags.append('ALARM')
        if d['fault']: flags.append('FAULT')
        if d['disabled']: flags.append('DISABLED')
        r[5].text = ', '.join(flags) if flags else 'OK'
        r[6].text = ''

    # 7. SYSTEM ANALYSIS & PERFORMANCE
    doc.add_heading('7. SYSTEM ANALYSIS & PERFORMANCE', level=1)
    doc.add_paragraph('Interpretation of system behavior:')
    doc.add_paragraph('- Network faults are the dominant fault type and require engineering assessment.')
    doc.add_paragraph('- Zone 1 and Zone 5 show increased activity and should be inspected.')
    doc.add_paragraph('- Replace devices flagged as persistent faults (see device table).')

    # 8. RECOMMENDATIONS & ACTION PLAN
    doc.add_heading('8. RECOMMENDATIONS & ACTION PLAN', level=1)
    doc.add_paragraph('Immediate Actions (0-7 days):')
    doc.add_paragraph('- Replace Device 180 (Loop 1) if flagged as persistent fault.')
    doc.add_paragraph('- Inspect short-circuit conditions on devices with repeated faults.')
    doc.add_paragraph('Preventive Actions (30 days):')
    doc.add_paragraph('- Clean sensitive optical detectors in high-dust areas.')
    doc.add_paragraph('- Run loop integrity tests and network diagnostics.')

    # --------------------------------------------------------
    # STEP 1: Save DOCX document
    # --------------------------------------------------------
    try:
        doc.save(OUTPUT_DOCX)
        print("✅ DOCX Report saved to:", OUTPUT_DOCX)
    except Exception as e:
        print(f"❌ ERROR: Could not save DOCX report to {OUTPUT_DOCX}. Error: {e}")
        return
